import streamlit as st
import pandas as pd
import joblib
import matplotlib.pyplot as plt
from io import BytesIO
from fpdf import FPDF
from docx import Document

# ---------------- Custom CSS ----------------
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Poppins:wght@400;600;700&display=swap');

body {
    background: linear-gradient(120deg, #d7ffd9, #ffffff);
    font-family: 'Poppins', sans-serif;
}

.main-title {
    text-align:center;
    font-size:46px;
    font-weight:700;
    color:#1B5E20;
    text-shadow: 2px 2px 10px rgba(0,0,0,0.1);
}

.version-tag {
    text-align:center;
    font-size:18px;
    color:#2e7d32;
    margin-top:-15px;
    margin-bottom:10px;
}

.stButton>button {
    background: linear-gradient(to right, #4caf50, #2e7d32) !important;
    color: white !important;
    font-size: 20px;
    font-weight:600;
    border-radius:12px;
    padding:12px 25px;
    transition: transform 0.2s;
}
.stButton>button:hover {
    transform: scale(1.05);
    box-shadow: 0 8px 20px rgba(0,0,0,0.3);
}

.result-card {
    margin-top:20px;
    padding:25px;
    border-radius:15px;
    text-align:center;
    font-size:24px;
    font-weight:600;
    box-shadow: 0px 10px 25px rgba(0,0,0,0.2);
    transition: transform 0.2s;
}
.result-card:hover {
    transform: scale(1.02);
}

.footer {
    text-align:center;
    margin-top:40px;
    color:#2e7d32;
    font-size:16px;
    font-weight:600;
}

.sidebar .sidebar-content {
    background: linear-gradient(to bottom, #e8f5e9, #ffffff);
    border-radius:15px;
    padding:15px;
}
</style>
""", unsafe_allow_html=True)

# ---------------- App Title ----------------
st.markdown("<h1 class='main-title'>Diabetes Prediction System</h1>", unsafe_allow_html=True)
st.markdown("<p class='version-tag'>Version 2.6 • Designed by Abhijit Shete</p>", unsafe_allow_html=True)
st.markdown("---")

# ---------------- Sidebar Inputs ----------------
st.sidebar.header("🧾 Customer Details")
customer_name = st.sidebar.text_input("👤 Customer Name")
customer_address = st.sidebar.text_area("🏠 Address")
customer_mobile = st.sidebar.text_input("📞 Mobile Number")

st.sidebar.markdown("---")
st.sidebar.header("📌 Enter Your Health Details or Upload CSV")

file = st.sidebar.file_uploader("Upload CSV to auto-fill values", type=["csv"])

pregnancies = glucose = bp = skin = insulin = bmi = dpf = age = None

if file is not None:
    df = pd.read_csv(file)
    pregnancies = df["Pregnancies"].iloc[0]
    glucose = df["Glucose"].iloc[0]
    bp = df["BloodPressure"].iloc[0]
    skin = df["SkinThickness"].iloc[0]
    insulin = df["Insulin"].iloc[0]
    bmi = df["BMI"].iloc[0]
    dpf = df["DPF"].iloc[0]
    age = df["Age"].iloc[0]
    st.sidebar.success("✔ CSV Loaded Successfully!")

if pregnancies is None:
    pregnancies = st.sidebar.number_input("👶 Pregnancies", 0, 20, 0)
    glucose = st.sidebar.number_input("🍬 Glucose Level", 0, 200, 100)
    bp = st.sidebar.number_input("💓 Blood Pressure", 0, 140, 70)
    skin = st.sidebar.number_input("📏 Skin Thickness", 0, 100, 20)
    insulin = st.sidebar.number_input("💉 Insulin", 0, 900, 100)
    bmi = st.sidebar.number_input("⚖ BMI", 0.0, 70.0, 25.0)
    dpf = st.sidebar.number_input("🧬 Diabetes Pedigree Function", 0.0, 3.0, 0.5)
    age = st.sidebar.number_input("🎂 Age", 1, 120, 30)

# ---------------- Load model ----------------
model = joblib.load("model.pkl")
scaler = joblib.load("scaler.pkl")

# ---------------- Report Generators ----------------
def generate_docx(result, values):
    doc = Document()
    doc.add_heading("Diabetes Prediction Report", level=1)
    doc.add_paragraph(f"Prediction Result: {result}")
    doc.add_heading("Input Health Metrics", level=2)
    for key, value in values.items():
        doc.add_paragraph(f"{key}: {value}")
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

# ---------------- PDF Function (DEPLOYMENT-SAFE) ----------------
def generate_pdf(result, values, customer):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", 'B', 14)

    result_safe = result.replace("—", "-")

    pdf.cell(200, 10, txt="Diabetes Prediction Report", ln=True, align='C')
    pdf.ln(5)

    pdf.set_font("Arial", 'B', 12)
    pdf.cell(200, 8, txt="Customer Details", ln=True)
    pdf.set_font("Arial", size=11)
    pdf.cell(200, 7, txt=f"Name: {customer['name']}", ln=True)
    pdf.cell(200, 7, txt=f"Mobile No: {customer['mobile']}", ln=True)
    pdf.multi_cell(0, 7, txt=f"Address: {customer['address']}")

    pdf.ln(4)
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt=f"Prediction Result: {result_safe}", ln=True)
    pdf.ln(3)

    pdf.set_font("Arial", 'B', 12)
    pdf.cell(200, 10, txt="Input Health Metrics:", ln=True)
    pdf.set_font("Arial", size=11)
    for key, value in values.items():
        pdf.cell(200, 8, txt=f"{key}: {value}", ln=True)

    pdf_output = pdf.output(dest="S")

    # ✅ FINAL DEPLOYMENT-SAFE FIX
    if isinstance(pdf_output, str):
        return pdf_output.encode("latin1")
    else:
        return bytes(pdf_output)

# ---------------- Prediction ----------------
predict = st.button("🚀 Predict")
if predict:
    data = pd.DataFrame([[pregnancies, glucose, bp, skin, insulin, bmi, dpf, age]],
                        columns=["Pregnancies","Glucose","BloodPressure","SkinThickness",
                                 "Insulin","BMI","DiabetesPedigreeFunction","Age"])

    data_scaled = scaler.transform(data)
    prediction = model.predict(data_scaled)
    prob = model.predict_proba(data_scaled)[0][1]

    result_text = "High Risk of Diabetes!" if prediction[0] == 1 else "Low Risk - You are Safe!"

    if prediction[0] == 1:
        st.markdown(
            "<div class='result-card' style='background: linear-gradient(to right, #ff8a80, #d32f2f); color:white;'>⚠ High Risk of Diabetes!</div>",
            unsafe_allow_html=True)
    else:
        st.markdown(
            "<div class='result-card' style='background: linear-gradient(to right, #a5d6a7, #2e7d32); color:white;'>✅ Low Risk - You are Safe!</div>",
            unsafe_allow_html=True)

    st.info(f"Prediction Probability: {prob*100:.2f}%")

    st.subheader("📈 Your Health Metrics (Line Chart View)")
    labels = ["Pregnancies", "Glucose", "BloodPressure", "SkinThickness",
              "Insulin", "BMI", "DPF", "Age"]
    values_list = [pregnancies, glucose, bp, skin, insulin, bmi, dpf, age]

    fig, ax = plt.subplots(figsize=(10,5))
    ax.plot(labels, values_list, marker="o", linewidth=2, color="green")
    ax.grid(True, alpha=0.2)
    plt.xticks(rotation=30)
    st.pyplot(fig)

    values_dict = {
        "Pregnancies": pregnancies,
        "Glucose": glucose,
        "Blood Pressure": bp,
        "Skin Thickness": skin,
        "Insulin": insulin,
        "BMI": bmi,
        "Diabetes Pedigree": dpf,
        "Age": age
    }

    customer_details = {
        "name": customer_name or "N/A",
        "address": customer_address or "N/A",
        "mobile": customer_mobile or "N/A"
    }

    docx_file = generate_docx(result_text, values_dict)
    st.download_button(
        label="📄 Download DOCX Report",
        data=docx_file,
        file_name="diabetes_report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    pdf_file = generate_pdf(result_text, values_dict, customer_details)
    st.download_button(
        label="📄 Download Premium PDF Report",
        data=pdf_file,
        file_name="diabetes_report.pdf",
        mime="application/pdf"
    )

# ---------------- About ----------------
with st.expander("ℹ️ About this Application"):
    st.write("""
    This app predicts diabetes risk using machine learning.
    """)

st.markdown("<p class='footer'>© 2025 | Designed by Abhijit Shete</p>", unsafe_allow_html=True)
